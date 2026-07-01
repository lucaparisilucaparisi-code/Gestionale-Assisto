"""
Route di export/report (Excel, Word) - Blueprint separato da app.py.

Contiene gli export: mensile completo, annuale, municipale,
dipartimentale e la relazione Word.
"""

import io
from datetime import datetime

from flask import Blueprint, request, send_file
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import config
import database as db

logger = config.setup_logging()
MESI_NOME = config.MESI_NOME
MESI_SCOLASTICI = config.MESI_SCOLASTICI

export_bp = Blueprint('export', __name__)


def decimal_to_sessagesimal(decimal_hours):
    """
    Converte ore decimali in formato HH:MM
    Es: 1.50 → "1:30", 2.25 → "2:15", 0.75 → "0:45"
    """
    if decimal_hours is None or decimal_hours == 0:
        return "0:00"

    hours = int(decimal_hours)
    minutes = round((decimal_hours - hours) * 60)

    # Gestisci arrotondamento
    if minutes == 60:
        hours += 1
        minutes = 0

    return f"{hours}:{minutes:02d}"


# Mappa per ordinamento cronologico dei mesi di lista attesa (anno scolastico)
_LISTA_ATTESA_MESI_ORDER = {
    'Settembre': 9, 'Ottobre': 10, 'Novembre': 11, 'Dicembre': 12,
    'Gennaio': 1, 'Febbraio': 2, 'Marzo': 3, 'Aprile': 4, 'Maggio': 5, 'Giugno': 6,
}
_LISTA_ATTESA_MESI_ABBR = {
    'Settembre': 'Set', 'Ottobre': 'Ott', 'Novembre': 'Nov', 'Dicembre': 'Dic',
    'Gennaio': 'Gen', 'Febbraio': 'Feb', 'Marzo': 'Mar', 'Aprile': 'Apr',
    'Maggio': 'Mag', 'Giugno': 'Giu',
}
# Ordine cronologico nell'anno scolastico (Set -> Giu)
_LISTA_ATTESA_SORT_INDEX = ['Settembre', 'Ottobre', 'Novembre', 'Dicembre',
                            'Gennaio', 'Febbraio', 'Marzo', 'Aprile', 'Maggio', 'Giugno']


def get_liste_attesa_ordinate(dati, anno_report, mese_report):
    """Ritorna i valori distinti di lista_attesa presenti nei dati, ordinati
    cronologicamente per anno scolastico, con etichetta 'Lista Mes YYYY'.

    Returns: list of dict { 'valore': 'Settembre', 'label': 'Lista Set 2025', 'anno': 2025 }
    """
    # Determina anno di inizio dell'anno scolastico
    if mese_report >= 9:
        anno_inizio_as = anno_report
    else:
        anno_inizio_as = anno_report - 1

    valori = set()
    for d in dati:
        la = (d.get('lista_attesa') or '').strip()
        if la:
            valori.add(la)

    risultato = []
    for v in valori:
        if v in _LISTA_ATTESA_SORT_INDEX:
            idx = _LISTA_ATTESA_SORT_INDEX.index(v)
            anno_v = anno_inizio_as if idx < 4 else anno_inizio_as + 1
            abbr = _LISTA_ATTESA_MESI_ABBR[v]
            label = f"Lista {abbr} {anno_v}"
            risultato.append({'valore': v, 'label': label, 'anno': anno_v, 'sort_idx': idx})
        else:
            # Valore non riconosciuto: lo mettiamo in fondo con label grezza
            risultato.append({'valore': v, 'label': f"Lista {v}", 'anno': 9999, 'sort_idx': 99})

    risultato.sort(key=lambda x: x['sort_idx'])
    return risultato


# ==================== BRAND / STILI REPORT ====================

# Palette report (coerente con il brand)
REPORT_PRIMARY = '#4F46E5'        # Indigo
REPORT_PRIMARY_DARK = '#3730A3'   # Indigo scuro
REPORT_ACCENT = '#0EA5E9'         # Sky blue
REPORT_DARK = '#1E293B'           # Slate scuro (per testi e totali)
REPORT_MUTED = '#64748B'          # Slate medio
REPORT_LIGHT = '#F1F5F9'          # Slate chiaro (righe alternate)
REPORT_BORDER = '#CBD5E1'         # Bordo cella tenue
REPORT_SUCCESS = '#10B981'        # Verde


def get_excel_brand_styles(workbook):
    """Restituisce un dizionario di formati Excel branded e coerenti."""
    FONT = 'Calibri'

    return {
        'title': workbook.add_format({
            'bold': True, 'font_size': 18, 'font_name': FONT,
            'bg_color': REPORT_PRIMARY, 'font_color': 'white',
            'align': 'center', 'valign': 'vcenter',
        }),
        'subtitle': workbook.add_format({
            'bold': True, 'font_size': 11, 'font_name': FONT,
            'bg_color': REPORT_DARK, 'font_color': 'white',
            'align': 'center', 'valign': 'vcenter',
        }),
        'section': workbook.add_format({
            'bold': True, 'font_size': 12, 'font_name': FONT,
            'bg_color': REPORT_DARK, 'font_color': 'white',
            'align': 'left', 'valign': 'vcenter', 'indent': 1,
        }),
        'info': workbook.add_format({
            'italic': True, 'font_size': 9, 'font_name': FONT,
            'font_color': REPORT_MUTED, 'align': 'left', 'valign': 'vcenter',
        }),
        'header': workbook.add_format({
            'bold': True, 'font_size': 10, 'font_name': FONT,
            'bg_color': REPORT_PRIMARY, 'font_color': 'white',
            'align': 'center', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_PRIMARY_DARK,
            'text_wrap': True,
        }),
        'cell': workbook.add_format({
            'font_size': 10, 'font_name': FONT, 'align': 'left',
            'valign': 'vcenter', 'border': 1, 'border_color': REPORT_BORDER,
        }),
        'cell_alt': workbook.add_format({
            'font_size': 10, 'font_name': FONT, 'align': 'left',
            'valign': 'vcenter', 'border': 1, 'border_color': REPORT_BORDER,
            'bg_color': REPORT_LIGHT,
        }),
        'cell_center': workbook.add_format({
            'font_size': 10, 'font_name': FONT, 'align': 'center',
            'valign': 'vcenter', 'border': 1, 'border_color': REPORT_BORDER,
        }),
        'cell_center_alt': workbook.add_format({
            'font_size': 10, 'font_name': FONT, 'align': 'center',
            'valign': 'vcenter', 'border': 1, 'border_color': REPORT_BORDER,
            'bg_color': REPORT_LIGHT,
        }),
        'number': workbook.add_format({
            'font_size': 10, 'font_name': FONT, 'num_format': '#,##0.00',
            'align': 'right', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_BORDER,
        }),
        'number_alt': workbook.add_format({
            'font_size': 10, 'font_name': FONT, 'num_format': '#,##0.00',
            'align': 'right', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_BORDER,
            'bg_color': REPORT_LIGHT,
        }),
        'integer': workbook.add_format({
            'font_size': 10, 'font_name': FONT, 'num_format': '#,##0',
            'align': 'right', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_BORDER,
        }),
        'integer_alt': workbook.add_format({
            'font_size': 10, 'font_name': FONT, 'num_format': '#,##0',
            'align': 'right', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_BORDER,
            'bg_color': REPORT_LIGHT,
        }),
        'money': workbook.add_format({
            'font_size': 10, 'font_name': FONT, 'num_format': '€ #,##0.00',
            'align': 'right', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_BORDER,
        }),
        'money_alt': workbook.add_format({
            'font_size': 10, 'font_name': FONT, 'num_format': '€ #,##0.00',
            'align': 'right', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_BORDER,
            'bg_color': REPORT_LIGHT,
        }),
        'total_label': workbook.add_format({
            'bold': True, 'font_size': 10, 'font_name': FONT,
            'bg_color': REPORT_DARK, 'font_color': 'white',
            'align': 'left', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_DARK,
            'indent': 1,
        }),
        'total_cell': workbook.add_format({
            'bold': True, 'font_size': 10, 'font_name': FONT,
            'bg_color': REPORT_DARK, 'font_color': 'white',
            'align': 'center', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_DARK,
        }),
        'total_number': workbook.add_format({
            'bold': True, 'font_size': 10, 'font_name': FONT,
            'bg_color': REPORT_DARK, 'font_color': 'white',
            'num_format': '#,##0.00', 'align': 'right', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_DARK,
        }),
        'total_integer': workbook.add_format({
            'bold': True, 'font_size': 10, 'font_name': FONT,
            'bg_color': REPORT_DARK, 'font_color': 'white',
            'num_format': '#,##0', 'align': 'right', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_DARK,
        }),
        'total_money': workbook.add_format({
            'bold': True, 'font_size': 10, 'font_name': FONT,
            'bg_color': REPORT_DARK, 'font_color': 'white',
            'num_format': '€ #,##0.00', 'align': 'right', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_DARK,
        }),
        'summary_label': workbook.add_format({
            'bold': True, 'font_size': 10, 'font_name': FONT,
            'bg_color': REPORT_LIGHT, 'font_color': REPORT_DARK,
            'align': 'right', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_BORDER,
        }),
        'summary_money': workbook.add_format({
            'font_size': 10, 'font_name': FONT, 'num_format': '€ #,##0.00',
            'align': 'right', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_BORDER,
        }),
        'grand_label': workbook.add_format({
            'bold': True, 'font_size': 11, 'font_name': FONT,
            'bg_color': REPORT_PRIMARY, 'font_color': 'white',
            'align': 'right', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_PRIMARY_DARK,
        }),
        'grand_money': workbook.add_format({
            'bold': True, 'font_size': 11, 'font_name': FONT,
            'bg_color': REPORT_PRIMARY, 'font_color': 'white',
            'num_format': '€ #,##0.00', 'align': 'right', 'valign': 'vcenter',
            'border': 1, 'border_color': REPORT_PRIMARY_DARK,
        }),
    }


def _style_word_cell_bg(cell, hex_color):
    """Imposta il colore di sfondo di una cella Word (hex senza '#')."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)


def style_word_table_header(table, bg_color='4F46E5', font_color='FFFFFF'):
    """Applica stile professionale alla prima riga (header) di una tabella Word."""
    if not table.rows:
        return
    hdr_row = table.rows[0]
    for cell in hdr_row.cells:
        _style_word_cell_bg(cell, bg_color)
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(font_color)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)


def style_word_table_alternating_rows(table, color_light='FFFFFF', color_alt='F1F5F9'):
    """Applica righe alternate alla tabella Word (skip header)."""
    for i, row in enumerate(table.rows):
        if i == 0:
            continue
        bg = color_alt if (i % 2 == 0) else color_light
        for cell in row.cells:
            _style_word_cell_bg(cell, bg)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = 'Calibri'
                    run.font.size = Pt(10)


def style_word_table_total_row(table, bg_color='1E293B', font_color='FFFFFF'):
    """Applica stile alla riga di totale (ultima riga) di una tabella Word."""
    if not table.rows:
        return
    last_row = table.rows[-1]
    for cell in last_row.cells:
        _style_word_cell_bg(cell, bg_color)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor.from_string(font_color)
                run.font.name = 'Calibri'
                run.font.size = Pt(10)


def add_word_page_number_footer(doc):
    """Aggiunge un footer con numero pagina al documento Word."""
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = p.add_run('Pagina ')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string('64748B')

    # PAGE field
    fld_char_begin = OxmlElement('w:fldChar')
    fld_char_begin.set(qn('w:fldCharType'), 'begin')
    instr_text = OxmlElement('w:instrText')
    instr_text.set(qn('xml:space'), 'preserve')
    instr_text.text = 'PAGE'
    fld_char_end = OxmlElement('w:fldChar')
    fld_char_end.set(qn('w:fldCharType'), 'end')

    page_run = p.add_run()
    page_run.font.size = Pt(9)
    page_run.font.color.rgb = RGBColor.from_string('64748B')
    page_run._r.append(fld_char_begin)
    page_run._r.append(instr_text)
    page_run._r.append(fld_char_end)

    run2 = p.add_run(' di ')
    run2.font.size = Pt(9)
    run2.font.color.rgb = RGBColor.from_string('64748B')

    # NUMPAGES field
    fld_char_begin2 = OxmlElement('w:fldChar')
    fld_char_begin2.set(qn('w:fldCharType'), 'begin')
    instr_text2 = OxmlElement('w:instrText')
    instr_text2.set(qn('xml:space'), 'preserve')
    instr_text2.text = 'NUMPAGES'
    fld_char_end2 = OxmlElement('w:fldChar')
    fld_char_end2.set(qn('w:fldCharType'), 'end')

    numpages_run = p.add_run()
    numpages_run.font.size = Pt(9)
    numpages_run.font.color.rgb = RGBColor.from_string('64748B')
    numpages_run._r.append(fld_char_begin2)
    numpages_run._r.append(instr_text2)
    numpages_run._r.append(fld_char_end2)



# ==================== EXPORT PREMIUM ====================

@export_bp.route('/api/export/excel/<int:anno>/<int:mese>')
def api_export_excel(anno, mese):
    """Esporta rendicontazione in Excel - Versione Premium"""
    commessa = request.args.get('commessa')
    privacy = request.args.get('privacy', 'false').lower() == 'true'

    dati = db.get_rendicontazione_completa(anno, mese, commessa)
    totali_scuola = db.get_totali_per_scuola(anno, mese, commessa)

    # Calcola totali ore
    ore_totali_60 = sum(d['ore_lavorate_60'] or 0 for d in dati)
    ore_totali_100 = sum(d['ore_lavorate_100'] or 0 for d in dati)
    ore_previste = sum(d['media_con_assenza_60'] or 0 for d in dati)

    # Calcolo fatturazione corretto (sul totale, non somma di arrotondamenti)
    imponibile_totale, iva_totale, totale_lordo = config.calcola_fatturazione(ore_totali_100)

    # Calcola statistiche avanzate
    totale_generale = {
        'num_utenti': len(dati),
        'ore_lavorate_60': ore_totali_60,
        'ore_lavorate_100': ore_totali_100,
        'ore_previste': ore_previste,
        'imponibile_100': imponibile_totale,
        'iva_100': iva_totale,
        'totale_100': totale_lordo,
        'pasti': sum(d['pasti'] or 0 for d in dati),
        'credito_debito': sum(d['credito_debito'] or 0 for d in dati)
    }

    # Calcola percentuale completamento
    if totale_generale['ore_previste'] > 0:
        perc_completamento = (totale_generale['ore_lavorate_60'] / totale_generale['ore_previste']) * 100
    else:
        perc_completamento = 0

    # Determina anno scolastico
    if mese >= 9:
        anno_scolastico = f"{anno}/{anno+1}"
    else:
        anno_scolastico = f"{anno-1}/{anno}"

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        workbook = writer.book

        # ========== DEFINIZIONE FORMATI PREMIUM ==========
        # Colori brand
        # Palette allineata a quella branded condivisa (coerenza tra i report)
        PRIMARY_COLOR = REPORT_PRIMARY
        PRIMARY_DARK = REPORT_PRIMARY_DARK
        SUCCESS_COLOR = REPORT_SUCCESS
        DANGER_COLOR = config.DANGER_COLOR
        DARK_COLOR = REPORT_DARK

        # Formato titolo principale
        title_fmt = workbook.add_format({
            'bold': True,
            'font_size': 24,
            'font_color': DARK_COLOR,
            'align': 'left',
            'valign': 'vcenter'
        })

        # Formato sottotitolo
        subtitle_fmt = workbook.add_format({
            'font_size': 12,
            'font_color': '#64748B',
            'align': 'left'
        })

        # Formato header tabella
        header_fmt = workbook.add_format({
            'bold': True,
            'font_size': 10,
            'font_color': 'white',
            'bg_color': PRIMARY_COLOR,
            'align': 'center',
            'valign': 'vcenter',
            'border': 1,
            'border_color': PRIMARY_DARK,
            'text_wrap': True
        })

        # Formato cella normale
        cell_fmt = workbook.add_format({
            'font_size': 10,
            'align': 'left',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1'
        })

        # Formato numeri
        number_fmt = workbook.add_format({
            'font_size': 10,
            'num_format': '#,##0.00',
            'align': 'right',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1'
        })

        # Formato valuta
        money_fmt = workbook.add_format({
            'font_size': 10,
            'num_format': '€ #,##0.00',
            'align': 'right',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1',
            'font_color': SUCCESS_COLOR
        })

        # Formato credito (positivo)
        credit_fmt = workbook.add_format({
            'font_size': 10,
            'num_format': '#,##0.00',
            'align': 'right',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1',
            'font_color': SUCCESS_COLOR,
            'bold': True
        })

        # Formato debito (negativo)
        debit_fmt = workbook.add_format({
            'font_size': 10,
            'num_format': '#,##0.00',
            'align': 'right',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1',
            'font_color': DANGER_COLOR,
            'bold': True
        })

        # Formato totale riga
        total_fmt = workbook.add_format({
            'bold': True,
            'font_size': 10,
            'bg_color': '#EEF2FF',
            'align': 'right',
            'valign': 'vcenter',
            'border': 2,
            'border_color': PRIMARY_COLOR
        })

        # Formato totale valuta
        total_money_fmt = workbook.add_format({
            'bold': True,
            'font_size': 10,
            'num_format': '€ #,##0.00',
            'bg_color': '#EEF2FF',
            'align': 'right',
            'valign': 'vcenter',
            'border': 2,
            'border_color': PRIMARY_COLOR,
            'font_color': PRIMARY_DARK
        })

        # Formato KPI
        kpi_label_fmt = workbook.add_format({
            'font_size': 10,
            'font_color': '#64748B',
            'align': 'left',
            'valign': 'vcenter'
        })

        kpi_value_fmt = workbook.add_format({
            'bold': True,
            'font_size': 14,
            'font_color': DARK_COLOR,
            'align': 'left',
            'valign': 'vcenter'
        })

        kpi_money_fmt = workbook.add_format({
            'bold': True,
            'font_size': 14,
            'font_color': SUCCESS_COLOR,
            'num_format': '€ #,##0.00',
            'align': 'left',
            'valign': 'vcenter'
        })

        # ========== FOGLIO 1: RIEPILOGO ESECUTIVO ==========
        ws_summary = workbook.add_worksheet('Riepilogo')

        # Titolo report
        ws_summary.set_row(0, 40)
        ws_summary.write('A1', f'RENDICONTAZIONE OEPAC - {MESI_NOME[mese]} {anno}', title_fmt)
        ws_summary.write('A2', f'Anno Scolastico {anno_scolastico}', subtitle_fmt)
        if commessa:
            ws_summary.write('A3', f'Commessa: {commessa}', subtitle_fmt)

        ws_summary.write('A5', f'Generato il: {datetime.now().strftime("%d/%m/%Y alle %H:%M")}', subtitle_fmt)

        # KPI Cards
        ws_summary.write('A7', 'INDICATORI CHIAVE DI PERFORMANCE', workbook.add_format({
            'bold': True, 'font_size': 14, 'font_color': DARK_COLOR
        }))

        # Converti ore 60' in formato HH:MM per visualizzazione
        ore_60_formatted = decimal_to_sessagesimal(totale_generale['ore_lavorate_60'])

        kpis = [
            ('Utenti Totali', totale_generale['num_utenti'], None),
            ('Ore Lavorate (60\')', ore_60_formatted, None),
            ('Ore Previste (-11%)', round(totale_generale['ore_previste'], 2), None),
            ('% Completamento', f'{perc_completamento:.1f}%', None),
            ('Imponibile Totale', totale_generale['imponibile_100'], 'money'),
            ('IVA 5%', totale_generale['iva_100'], 'money'),
            ('Totale Lordo', totale_generale['totale_100'], 'money'),
            ('Credito/Debito Ore', round(totale_generale['credito_debito'], 2), None),
            ('Pasti Totali', totale_generale['pasti'], None),
        ]

        row = 8
        col = 0
        for i, (label, value, fmt_type) in enumerate(kpis):
            if i > 0 and i % 3 == 0:
                row += 3
                col = 0

            ws_summary.write(row, col, label, kpi_label_fmt)
            if fmt_type == 'money':
                ws_summary.write(row + 1, col, value, kpi_money_fmt)
            else:
                ws_summary.write(row + 1, col, value, kpi_value_fmt)
            col += 2

        # Riepilogo per scuola
        ws_summary.write(row + 5, 0, 'RIEPILOGO PER SCUOLA', workbook.add_format({
            'bold': True, 'font_size': 14, 'font_color': DARK_COLOR
        }))

        headers_scuola = ['Commessa', 'Scuola', 'Utenti', 'Ore (60\')', 'Ore (100\')', 'Imponibile', 'Totale']
        for c, h in enumerate(headers_scuola):
            ws_summary.write(row + 7, c, h, header_fmt)

        for i, t in enumerate(totali_scuola):
            r = row + 8 + i
            ws_summary.write(r, 0, t['commessa'], cell_fmt)
            ws_summary.write(r, 1, t['scuola'], cell_fmt)
            ws_summary.write(r, 2, t['num_utenti'], number_fmt)
            ws_summary.write(r, 3, decimal_to_sessagesimal(t['ore_lavorate_60']), cell_fmt)
            ws_summary.write(r, 4, t['ore_lavorate_100'], number_fmt)
            ws_summary.write(r, 5, t['imponibile_100'], money_fmt)
            ws_summary.write(r, 6, t['totale_100'], money_fmt)

        # Imposta larghezza colonne - B più larga per nomi scuole completi
        ws_summary.set_column('A:A', 20)
        ws_summary.set_column('B:B', 60)
        ws_summary.set_column('C:G', 15)

        # ========== FOGLIO 2: DETTAGLIO COMPLETO ==========
        ws_detail = workbook.add_worksheet('Dettaglio')

        # Header
        ws_detail.set_row(0, 30)
        ws_detail.write('A1', f'Dettaglio Rendicontazione - {MESI_NOME[mese]} {anno}', title_fmt)

        headers = [
            'Commessa', 'Scuola', 'Utente', 'A.C.', 'Monte Ore',
            'Media Mens.', 'Media -11%', 'Ore Lav. (60\')', 'Ore (100\')',
            'Imponibile', 'IVA 5%', 'Totale', 'Pasti', 'Cred/Deb', 'Lista Attesa'
        ]

        for c, h in enumerate(headers):
            ws_detail.write(2, c, h, header_fmt)

        # Dati
        for i, d in enumerate(dati):
            r = 3 + i
            utente = d['nome_puntato'] if privacy else f"{d['nome']} {d['cognome']}"

            ws_detail.write(r, 0, d['commessa'], cell_fmt)
            ws_detail.write(r, 1, d['scuola'], cell_fmt)
            ws_detail.write(r, 2, utente, cell_fmt)
            ws_detail.write(r, 3, d['nome_puntato'], cell_fmt)
            ws_detail.write(r, 4, d['monte_ore_settimanale'], number_fmt)
            ws_detail.write(r, 5, d['media_mensile_60'], number_fmt)
            ws_detail.write(r, 6, d['media_con_assenza_60'], number_fmt)
            ws_detail.write(r, 7, decimal_to_sessagesimal(d['ore_lavorate_60'] or 0), cell_fmt)
            ws_detail.write(r, 8, d['ore_lavorate_100'] or 0, number_fmt)
            ws_detail.write(r, 9, d['imponibile_100'] or 0, money_fmt)
            ws_detail.write(r, 10, d['iva_100'] or 0, money_fmt)
            ws_detail.write(r, 11, d['totale_100'] or 0, money_fmt)
            ws_detail.write(r, 12, d['pasti'] or 0, number_fmt)

            # Credito/Debito con colore condizionale
            cd = d['credito_debito'] or 0
            ws_detail.write(r, 13, cd, credit_fmt if cd >= 0 else debit_fmt)

            # Lista Attesa
            ws_detail.write(r, 14, d.get('lista_attesa') or '', cell_fmt)

        # Riga totali
        total_row = 3 + len(dati)
        ws_detail.write(total_row, 0, 'TOTALE', total_fmt)
        ws_detail.write(total_row, 7, decimal_to_sessagesimal(totale_generale['ore_lavorate_60']), total_fmt)
        ws_detail.write(total_row, 8, totale_generale['ore_lavorate_100'], total_fmt)
        ws_detail.write(total_row, 9, totale_generale['imponibile_100'], total_money_fmt)
        ws_detail.write(total_row, 10, totale_generale['iva_100'], total_money_fmt)
        ws_detail.write(total_row, 11, totale_generale['totale_100'], total_money_fmt)
        ws_detail.write(total_row, 12, totale_generale['pasti'], total_fmt)
        ws_detail.write(total_row, 13, totale_generale['credito_debito'], total_fmt)
        ws_detail.write(total_row, 14, '', total_fmt)

        # Larghezza colonne - B più larga per nomi scuole completi
        ws_detail.set_column('A:A', 12)
        ws_detail.set_column('B:B', 55)
        ws_detail.set_column('C:C', 25)
        ws_detail.set_column('D:D', 8)
        ws_detail.set_column('E:N', 12)
        ws_detail.set_column('O:O', 12)

        # Freeze panes
        ws_detail.freeze_panes(3, 0)

        # ========== FOGLIO 3: ANALISI STATISTICHE ==========
        ws_stats = workbook.add_worksheet('Statistiche')

        ws_stats.write('A1', 'ANALISI STATISTICHE', title_fmt)
        ws_stats.write('A3', 'Distribuzione Ore per Commessa', workbook.add_format({
            'bold': True, 'font_size': 12
        }))

        # Raggruppa per commessa
        commesse_stats = {}
        for d in dati:
            c = d['commessa']
            if c not in commesse_stats:
                commesse_stats[c] = {'ore': 0, 'utenti': 0, 'totale': 0}
            commesse_stats[c]['ore'] += d['ore_lavorate_60'] or 0
            commesse_stats[c]['utenti'] += 1
            commesse_stats[c]['totale'] += d['totale_100'] or 0

        ws_stats.write(4, 0, 'Commessa', header_fmt)
        ws_stats.write(4, 1, 'Utenti', header_fmt)
        ws_stats.write(4, 2, 'Ore Totali', header_fmt)
        ws_stats.write(4, 3, 'Fatturato', header_fmt)

        row = 5
        for c, stats in commesse_stats.items():
            ws_stats.write(row, 0, c, cell_fmt)
            ws_stats.write(row, 1, stats['utenti'], number_fmt)
            ws_stats.write(row, 2, stats['ore'], number_fmt)
            ws_stats.write(row, 3, stats['totale'], money_fmt)
            row += 1

        # Grafico a torta per distribuzione ore
        if len(commesse_stats) > 0:
            chart = workbook.add_chart({'type': 'pie'})
            chart.add_series({
                'name': 'Distribuzione Ore',
                'categories': f'=Statistiche!$A$6:$A${5 + len(commesse_stats)}',
                'values': f'=Statistiche!$C$6:$C${5 + len(commesse_stats)}',
                'data_labels': {'percentage': True, 'category': True}
            })
            chart.set_title({'name': 'Distribuzione Ore per Commessa'})
            chart.set_style(10)
            ws_stats.insert_chart('F3', chart, {'x_scale': 1.2, 'y_scale': 1.2})

        # Grafico a barre per fatturato
        if len(commesse_stats) > 0:
            chart2 = workbook.add_chart({'type': 'column'})
            chart2.add_series({
                'name': 'Fatturato',
                'categories': f'=Statistiche!$A$6:$A${5 + len(commesse_stats)}',
                'values': f'=Statistiche!$D$6:$D${5 + len(commesse_stats)}',
                'fill': {'color': PRIMARY_COLOR}
            })
            chart2.set_title({'name': 'Fatturato per Commessa'})
            chart2.set_style(10)
            chart2.set_y_axis({'num_format': '€ #,##0'})
            ws_stats.insert_chart('F18', chart2, {'x_scale': 1.2, 'y_scale': 1.2})

        ws_stats.set_column('A:D', 15)

        # ========== FOGLIO 4: DETTAGLIO PER SCUOLA (RAGGRUPPATO) ==========
        ws_scuola = workbook.add_worksheet('Dettaglio per Scuola')

        # Formato per header scuola
        scuola_header_fmt = workbook.add_format({
            'bold': True,
            'font_size': 11,
            'font_color': 'white',
            'bg_color': '#5B5FC7',
            'align': 'left',
            'valign': 'vcenter',
            'border': 1
        })

        # Formato celle utente
        utente_cell_fmt = workbook.add_format({
            'font_size': 9,
            'align': 'left',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1',
            'indent': 1
        })

        utente_number_fmt = workbook.add_format({
            'font_size': 9,
            'num_format': '#,##0.00',
            'align': 'right',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1'
        })

        utente_money_fmt = workbook.add_format({
            'font_size': 9,
            'num_format': '€ #,##0.00',
            'align': 'right',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1'
        })

        # Header del foglio
        ws_scuola.set_row(0, 30)
        ws_scuola.write('A1', f'Dettaglio per Scuola - {MESI_NOME[mese]} {anno}', title_fmt)

        # Raggruppa dati per scuola
        scuole_dict = {}
        for d in dati:
            scuola = d['scuola']
            if scuola not in scuole_dict:
                scuole_dict[scuola] = []
            scuole_dict[scuola].append(d)

        # Headers colonne dati utente
        detail_headers = ['Nome Puntato', 'Monte Ore', 'Media Mens.', 'Media -11%', 'Ore Lav. (60\')',
                          'Ore (100\')', 'Imponibile', 'IVA 5%', 'Totale', 'Pasti', 'Cred/Deb', 'Lista Attesa']

        row = 3
        for scuola, utenti in sorted(scuole_dict.items()):
            # Riga header scuola (espandibile)
            ws_scuola.merge_range(row, 0, row, len(detail_headers), f'⊟ {scuola}', scuola_header_fmt)
            row += 1

            # Header colonne per questa scuola
            for col, h in enumerate(detail_headers):
                ws_scuola.write(row, col, h, header_fmt)
            row += 1

            # Dati utenti
            for u in utenti:
                nome_puntato = u['nome_puntato'] if privacy else f"{u['nome']} {u['cognome']}"
                ws_scuola.write(row, 0, nome_puntato, utente_cell_fmt)
                ws_scuola.write(row, 1, u['monte_ore_settimanale'], utente_number_fmt)
                ws_scuola.write(row, 2, u['media_mensile_60'] or 0, utente_number_fmt)
                ws_scuola.write(row, 3, u['media_con_assenza_60'] or 0, utente_number_fmt)
                ws_scuola.write(row, 4, decimal_to_sessagesimal(u['ore_lavorate_60'] or 0), utente_cell_fmt)
                ws_scuola.write(row, 5, u['ore_lavorate_100'] or 0, utente_number_fmt)
                ws_scuola.write(row, 6, u['imponibile_100'] or 0, utente_money_fmt)
                ws_scuola.write(row, 7, u['iva_100'] or 0, utente_money_fmt)
                ws_scuola.write(row, 8, u['totale_100'] or 0, utente_money_fmt)
                ws_scuola.write(row, 9, u['pasti'] or 0, utente_number_fmt)
                ws_scuola.write(row, 10, u['credito_debito'] or 0, utente_number_fmt)
                ws_scuola.write(row, 11, u.get('lista_attesa') or '', utente_cell_fmt)
                row += 1

            # Riga vuota tra scuole
            row += 1

        # Larghezza colonne
        ws_scuola.set_column('A:A', 15)
        ws_scuola.set_column('B:L', 12)

    output.seek(0)

    filename = f"OEPAC_Rendicontazione_{MESI_NOME[mese]}_{anno}"
    if commessa:
        filename += f"_{commessa.replace(' ', '_')}"
    if privacy:
        filename += "_privacy"
    filename += ".xlsx"

    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename
    )


@export_bp.route('/api/export/annuale/<anno_scolastico>')
def api_export_annuale(anno_scolastico):
    """Esporta rendicontazione annuale - Versione Premium"""
    commessa = request.args.get('commessa')
    privacy = request.args.get('privacy', 'false').lower() == 'true'

    # Parse anno scolastico
    anni = anno_scolastico.split('-')
    anno_inizio = int(anni[0])
    anno_fine = int(anni[1])

    # Costanti
    TARIFFA = config.TARIFFA_ORARIA
    IVA_PERC = config.IVA_PERCENTUALE
    TASSO_ASSENZA = config.TASSO_ASSENZA

    # Raccogli tutti i dati dell'anno per calcoli aggregati
    tutti_dati_anno = {}  # {mese: dati}
    utenti_aggregati = {}  # {utente_id: {dati aggregati}}

    for mese in MESI_SCOLASTICI:
        anno = anno_inizio if mese >= 9 else anno_fine
        dati = db.get_rendicontazione_completa(anno, mese, commessa)
        tutti_dati_anno[mese] = {'anno': anno, 'dati': dati}

        # Aggrega per utente
        for d in dati:
            utente_key = d['utente_id']
            if utente_key not in utenti_aggregati:
                utenti_aggregati[utente_key] = {
                    'nome': d['nome'],
                    'cognome': d['cognome'],
                    'nome_puntato': d['nome_puntato'],
                    'scuola': d['scuola'],
                    'commessa': d['commessa'],
                    'monte_ore_settimanale': d['monte_ore_settimanale'],
                    'ore_erogate_totali': 0,
                    'monte_ore_previsto_totale': 0,  # Somma delle medie mensili -11%
                    'pasti_totali': 0,
                    'imponibile_totale': 0,
                    'mesi_attivi': 0
                }
            utenti_aggregati[utente_key]['ore_erogate_totali'] += d['ore_lavorate_60'] or 0
            utenti_aggregati[utente_key]['monte_ore_previsto_totale'] += d['media_con_assenza_60'] or 0
            utenti_aggregati[utente_key]['pasti_totali'] += d['pasti'] or 0
            utenti_aggregati[utente_key]['mesi_attivi'] += 1

    # Imponibile per-utente calcolato UNA volta sul totale ore (non somma di
    # arrotondamenti mensili): cosi' la somma della colonna quadra col totale annuale.
    for u in utenti_aggregati.values():
        u['imponibile_totale'] = config.calcola_fatturazione(u['ore_erogate_totali'])[0]

    # Calcola totali annuali
    totale_ore_60 = sum(
        sum(d['ore_lavorate_60'] or 0 for d in m['dati'])
        for m in tutti_dati_anno.values()
    )
    totale_ore_100 = sum(
        sum(d['ore_lavorate_100'] or 0 for d in m['dati'])
        for m in tutti_dati_anno.values()
    )
    totale_ore_previste = sum(
        sum(d['media_con_assenza_60'] or 0 for d in m['dati'])
        for m in tutti_dati_anno.values()
    )
    totale_pasti = sum(
        sum(d['pasti'] or 0 for d in m['dati'])
        for m in tutti_dati_anno.values()
    )
    imponibile_annuale, iva_annuale, totale_lordo_annuale = config.calcola_fatturazione(totale_ore_100)

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        workbook = writer.book

        # ========== DEFINIZIONE FORMATI PREMIUM ==========
        # Palette allineata a quella branded condivisa (coerenza tra i report)
        PRIMARY_COLOR = REPORT_PRIMARY
        PRIMARY_DARK = REPORT_PRIMARY_DARK
        SUCCESS_COLOR = REPORT_SUCCESS
        DANGER_COLOR = config.DANGER_COLOR
        DARK_COLOR = REPORT_DARK

        # Formato titolo principale
        title_fmt = workbook.add_format({
            'bold': True,
            'font_size': 24,
            'font_color': DARK_COLOR,
            'align': 'left',
            'valign': 'vcenter'
        })

        # Formato sottotitolo
        subtitle_fmt = workbook.add_format({
            'font_size': 12,
            'font_color': '#64748B',
            'align': 'left'
        })

        # Formato sezione
        section_fmt = workbook.add_format({
            'bold': True,
            'font_size': 14,
            'font_color': DARK_COLOR,
            'bottom': 2,
            'bottom_color': PRIMARY_COLOR
        })

        # Formato header tabella
        header_fmt = workbook.add_format({
            'bold': True,
            'font_size': 10,
            'font_color': 'white',
            'bg_color': PRIMARY_COLOR,
            'align': 'center',
            'valign': 'vcenter',
            'border': 1,
            'border_color': PRIMARY_DARK,
            'text_wrap': True
        })

        # Formato header alternativo (grigio)
        workbook.add_format({
            'bold': True,
            'font_size': 10,
            'font_color': 'white',
            'bg_color': '#475569',
            'align': 'center',
            'valign': 'vcenter',
            'border': 1,
            'text_wrap': True
        })

        # Formato cella normale
        cell_fmt = workbook.add_format({
            'font_size': 10,
            'align': 'left',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1'
        })

        # Formato cella alternata (zebra)
        cell_alt_fmt = workbook.add_format({
            'font_size': 10,
            'align': 'left',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1',
            'bg_color': '#F1F5F9'
        })

        # Formato numeri
        number_fmt = workbook.add_format({
            'font_size': 10,
            'num_format': '#,##0.00',
            'align': 'right',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1'
        })

        number_alt_fmt = workbook.add_format({
            'font_size': 10,
            'num_format': '#,##0.00',
            'align': 'right',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1',
            'bg_color': '#F1F5F9'
        })

        # Formato valuta
        money_fmt = workbook.add_format({
            'font_size': 10,
            'num_format': '€ #,##0.00',
            'align': 'right',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1',
            'font_color': SUCCESS_COLOR
        })

        money_alt_fmt = workbook.add_format({
            'font_size': 10,
            'num_format': '€ #,##0.00',
            'align': 'right',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1',
            'font_color': SUCCESS_COLOR,
            'bg_color': '#F1F5F9'
        })

        # Formato credito (positivo)
        credit_fmt = workbook.add_format({
            'font_size': 10,
            'num_format': '+#,##0.00;-#,##0.00',
            'align': 'right',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1',
            'font_color': SUCCESS_COLOR,
            'bold': True
        })

        credit_alt_fmt = workbook.add_format({
            'font_size': 10,
            'num_format': '+#,##0.00;-#,##0.00',
            'align': 'right',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1',
            'font_color': SUCCESS_COLOR,
            'bold': True,
            'bg_color': '#F1F5F9'
        })

        # Formato debito (negativo)
        debit_fmt = workbook.add_format({
            'font_size': 10,
            'num_format': '+#,##0.00;-#,##0.00',
            'align': 'right',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1',
            'font_color': DANGER_COLOR,
            'bold': True
        })

        debit_alt_fmt = workbook.add_format({
            'font_size': 10,
            'num_format': '+#,##0.00;-#,##0.00',
            'align': 'right',
            'valign': 'vcenter',
            'border': 1,
            'border_color': '#CBD5E1',
            'font_color': DANGER_COLOR,
            'bold': True,
            'bg_color': '#F1F5F9'
        })

        # Formato totale riga
        total_fmt = workbook.add_format({
            'bold': True,
            'font_size': 10,
            'bg_color': '#EEF2FF',
            'align': 'right',
            'valign': 'vcenter',
            'border': 2,
            'border_color': PRIMARY_COLOR
        })

        total_text_fmt = workbook.add_format({
            'bold': True,
            'font_size': 10,
            'bg_color': '#EEF2FF',
            'align': 'left',
            'valign': 'vcenter',
            'border': 2,
            'border_color': PRIMARY_COLOR
        })

        # Formato totale valuta
        total_money_fmt = workbook.add_format({
            'bold': True,
            'font_size': 10,
            'num_format': '€ #,##0.00',
            'bg_color': '#EEF2FF',
            'align': 'right',
            'valign': 'vcenter',
            'border': 2,
            'border_color': PRIMARY_COLOR,
            'font_color': PRIMARY_DARK
        })

        # Formato KPI
        kpi_label_fmt = workbook.add_format({
            'font_size': 10,
            'font_color': '#64748B',
            'align': 'left',
            'valign': 'vcenter'
        })

        kpi_value_fmt = workbook.add_format({
            'bold': True,
            'font_size': 16,
            'font_color': DARK_COLOR,
            'align': 'left',
            'valign': 'vcenter'
        })

        kpi_money_fmt = workbook.add_format({
            'bold': True,
            'font_size': 16,
            'font_color': SUCCESS_COLOR,
            'num_format': '€ #,##0.00',
            'align': 'left',
            'valign': 'vcenter'
        })

        workbook.add_format({
            'bg_color': '#F1F5F9',
            'border': 1,
            'border_color': '#CBD5E1'
        })

        # ========== FOGLIO 1: DASHBOARD ANNUALE ==========
        ws_dashboard = workbook.add_worksheet('Dashboard')

        # Titolo
        ws_dashboard.set_row(0, 45)
        ws_dashboard.merge_range('A1:H1', 'RIEPILOGO ANNUALE OEPAC', title_fmt)
        ws_dashboard.write('A2', f'Anno Scolastico {anno_scolastico}', subtitle_fmt)
        if commessa:
            ws_dashboard.write('A3', f'Commessa: {commessa}', subtitle_fmt)
        ws_dashboard.write('A4', f'Generato il: {datetime.now().strftime("%d/%m/%Y alle %H:%M")}', subtitle_fmt)

        # Sezione KPI principali
        ws_dashboard.write('A6', 'INDICATORI CHIAVE ANNUALI', section_fmt)

        # Calcola percentuale completamento
        perc_completamento = (totale_ore_60 / totale_ore_previste * 100) if totale_ore_previste > 0 else 0

        kpis = [
            ('Utenti Attivi', len(utenti_aggregati), None),
            ('Ore Erogate Totali', decimal_to_sessagesimal(totale_ore_60), None),
            ('Ore Previste (-11%)', decimal_to_sessagesimal(totale_ore_previste), None),
            ('Completamento', f'{perc_completamento:.1f}%', None),
            ('Imponibile Annuale', imponibile_annuale, 'money'),
            ('IVA 5%', iva_annuale, 'money'),
            ('Totale Lordo', totale_lordo_annuale, 'money'),
            ('Pasti Totali', totale_pasti, None),
        ]

        row = 7
        col = 0
        for i, (label, value, fmt_type) in enumerate(kpis):
            if i > 0 and i % 4 == 0:
                row += 3
                col = 0

            ws_dashboard.write(row, col, label, kpi_label_fmt)
            if fmt_type == 'money':
                ws_dashboard.write(row + 1, col, value, kpi_money_fmt)
            else:
                ws_dashboard.write(row + 1, col, value, kpi_value_fmt)
            col += 2

        # Sezione Riepilogo Mensile
        ws_dashboard.write(row + 5, 0, 'ANDAMENTO MENSILE', section_fmt)

        headers_mese = ['Mese', 'Utenti', 'Ore Erogate', 'Ore Previste', 'Completamento', 'Imponibile', 'Totale', 'Pasti']
        for c, h in enumerate(headers_mese):
            ws_dashboard.write(row + 7, c, h, header_fmt)

        riepilogo_row = row + 8
        totali_riepilogo = {'utenti': 0, 'ore': 0, 'previste': 0, 'imponibile': 0, 'totale': 0, 'pasti': 0}

        for mese in MESI_SCOLASTICI:
            anno = tutti_dati_anno[mese]['anno']
            dati = tutti_dati_anno[mese]['dati']

            ore_mese = sum(d['ore_lavorate_60'] or 0 for d in dati)
            ore_100_mese = sum(d['ore_lavorate_100'] or 0 for d in dati)
            ore_previste_mese = sum(d['media_con_assenza_60'] or 0 for d in dati)
            imponibile_mese = round(ore_100_mese * TARIFFA, 2)
            iva_mese = round(imponibile_mese * IVA_PERC, 2)
            totale_mese = round(imponibile_mese + iva_mese, 2)
            pasti_mese = sum(d['pasti'] or 0 for d in dati)
            perc_mese = (ore_mese / ore_previste_mese * 100) if ore_previste_mese > 0 else 0

            is_alt = (riepilogo_row - row - 8) % 2 == 1
            cf = cell_alt_fmt if is_alt else cell_fmt
            nf = number_alt_fmt if is_alt else number_fmt
            mf = money_alt_fmt if is_alt else money_fmt

            ws_dashboard.write(riepilogo_row, 0, f'{MESI_NOME[mese]} {anno}', cf)
            ws_dashboard.write(riepilogo_row, 1, len(dati), nf)
            ws_dashboard.write(riepilogo_row, 2, decimal_to_sessagesimal(ore_mese), cf)
            ws_dashboard.write(riepilogo_row, 3, decimal_to_sessagesimal(ore_previste_mese), cf)
            ws_dashboard.write(riepilogo_row, 4, f'{perc_mese:.1f}%', cf)
            ws_dashboard.write(riepilogo_row, 5, imponibile_mese, mf)
            ws_dashboard.write(riepilogo_row, 6, totale_mese, mf)
            ws_dashboard.write(riepilogo_row, 7, pasti_mese, nf)

            totali_riepilogo['ore'] += ore_mese
            totali_riepilogo['previste'] += ore_previste_mese
            totali_riepilogo['imponibile'] += imponibile_mese
            totali_riepilogo['totale'] += totale_mese
            totali_riepilogo['pasti'] += pasti_mese

            riepilogo_row += 1

        # Riga totali
        ws_dashboard.write(riepilogo_row, 0, 'TOTALE ANNUALE', total_text_fmt)
        ws_dashboard.write(riepilogo_row, 1, len(utenti_aggregati), total_fmt)
        ws_dashboard.write(riepilogo_row, 2, decimal_to_sessagesimal(totali_riepilogo['ore']), total_fmt)
        ws_dashboard.write(riepilogo_row, 3, decimal_to_sessagesimal(totali_riepilogo['previste']), total_fmt)
        perc_tot = (totali_riepilogo['ore'] / totali_riepilogo['previste'] * 100) if totali_riepilogo['previste'] > 0 else 0
        ws_dashboard.write(riepilogo_row, 4, f'{perc_tot:.1f}%', total_fmt)
        ws_dashboard.write(riepilogo_row, 5, totali_riepilogo['imponibile'], total_money_fmt)
        ws_dashboard.write(riepilogo_row, 6, totali_riepilogo['totale'], total_money_fmt)
        ws_dashboard.write(riepilogo_row, 7, totali_riepilogo['pasti'], total_fmt)

        # Larghezza colonne
        ws_dashboard.set_column('A:A', 18)
        ws_dashboard.set_column('B:H', 14)

        # ========== FOGLIO 2: RIEPILOGO UTENTI (NUOVO!) ==========
        ws_utenti = workbook.add_worksheet('Riepilogo Utenti')

        ws_utenti.set_row(0, 40)
        ws_utenti.merge_range('A1:J1', f'RIEPILOGO PER UTENTE - A.S. {anno_scolastico}', title_fmt)
        ws_utenti.write('A2', 'Vista aggregata delle ore erogate per ogni utente', subtitle_fmt)
        ws_utenti.write('A3', f'Monte ore con detrazione assenze previste: {int(TASSO_ASSENZA*100)}%', subtitle_fmt)

        headers_utenti = [
            'Utente', 'Scuola', 'Commessa', 'Monte Ore Sett.',
            'Mesi Attivi', 'Monte Ore Previsto', 'Ore Erogate',
            'Credito/Debito', 'Pasti', 'Imponibile'
        ]

        for c, h in enumerate(headers_utenti):
            ws_utenti.write(5, c, h, header_fmt)

        # Ordina utenti per cognome e nome
        utenti_sorted = sorted(
            utenti_aggregati.values(),
            key=lambda x: (x['cognome'].lower(), x['nome'].lower())
        )

        utente_row = 6
        for i, u in enumerate(utenti_sorted):
            is_alt = i % 2 == 1
            cf = cell_alt_fmt if is_alt else cell_fmt
            nf = number_alt_fmt if is_alt else number_fmt
            mf = money_alt_fmt if is_alt else money_fmt

            utente_nome = u['nome_puntato'] if privacy else f"{u['nome']} {u['cognome']}"
            credito_debito = u['monte_ore_previsto_totale'] - u['ore_erogate_totali']

            # Seleziona formato per credito/debito
            if credito_debito >= 0:
                cd_fmt = credit_alt_fmt if is_alt else credit_fmt
            else:
                cd_fmt = debit_alt_fmt if is_alt else debit_fmt

            ws_utenti.write(utente_row, 0, utente_nome, cf)
            ws_utenti.write(utente_row, 1, u['scuola'], cf)
            ws_utenti.write(utente_row, 2, u['commessa'], cf)
            ws_utenti.write(utente_row, 3, u['monte_ore_settimanale'], nf)
            ws_utenti.write(utente_row, 4, u['mesi_attivi'], nf)
            ws_utenti.write(utente_row, 5, decimal_to_sessagesimal(u['monte_ore_previsto_totale']), cf)
            ws_utenti.write(utente_row, 6, decimal_to_sessagesimal(u['ore_erogate_totali']), cf)
            ws_utenti.write(utente_row, 7, round(credito_debito, 2), cd_fmt)
            ws_utenti.write(utente_row, 8, u['pasti_totali'], nf)
            ws_utenti.write(utente_row, 9, u['imponibile_totale'], mf)

            utente_row += 1

        # Riga totali utenti
        tot_monte_previsto = sum(u['monte_ore_previsto_totale'] for u in utenti_sorted)
        tot_ore_erogate = sum(u['ore_erogate_totali'] for u in utenti_sorted)
        tot_credito_debito = tot_monte_previsto - tot_ore_erogate
        tot_pasti = sum(u['pasti_totali'] for u in utenti_sorted)
        tot_imponibile = sum(u['imponibile_totale'] for u in utenti_sorted)

        ws_utenti.write(utente_row, 0, 'TOTALE', total_text_fmt)
        ws_utenti.write(utente_row, 1, '', total_fmt)
        ws_utenti.write(utente_row, 2, '', total_fmt)
        ws_utenti.write(utente_row, 3, '', total_fmt)
        ws_utenti.write(utente_row, 4, len(utenti_sorted), total_fmt)
        ws_utenti.write(utente_row, 5, decimal_to_sessagesimal(tot_monte_previsto), total_fmt)
        ws_utenti.write(utente_row, 6, decimal_to_sessagesimal(tot_ore_erogate), total_fmt)
        ws_utenti.write(utente_row, 7, round(tot_credito_debito, 2), total_fmt)
        ws_utenti.write(utente_row, 8, tot_pasti, total_fmt)
        ws_utenti.write(utente_row, 9, tot_imponibile, total_money_fmt)

        # Larghezza colonne
        ws_utenti.set_column('A:A', 25)
        ws_utenti.set_column('B:B', 50)
        ws_utenti.set_column('C:C', 15)
        ws_utenti.set_column('D:J', 16)

        # Freeze header
        ws_utenti.freeze_panes(6, 0)

        # ========== FOGLI MENSILI (DETTAGLIO) ==========
        for mese in MESI_SCOLASTICI:
            anno = tutti_dati_anno[mese]['anno']
            dati = tutti_dati_anno[mese]['dati']

            sheet_name = f'{MESI_NOME[mese][:3]} {anno}'
            ws_mese = workbook.add_worksheet(sheet_name)

            # Titolo
            ws_mese.set_row(0, 30)
            ws_mese.merge_range('A1:L1', f'Dettaglio {MESI_NOME[mese]} {anno}', title_fmt)

            # Headers
            headers_mese = [
                'Commessa', 'Scuola', 'Utente', 'Monte Ore',
                'Media Mens.', 'Media -11%', 'Ore (60\')', 'Ore (100\')',
                'Imponibile', 'IVA 5%', 'Totale', 'Cred/Deb', 'Pasti'
            ]

            for c, h in enumerate(headers_mese):
                ws_mese.write(2, c, h, header_fmt)

            # Dati
            for i, d in enumerate(dati):
                r = 3 + i
                is_alt = i % 2 == 1
                cf = cell_alt_fmt if is_alt else cell_fmt
                nf = number_alt_fmt if is_alt else number_fmt
                mf = money_alt_fmt if is_alt else money_fmt

                utente = d['nome_puntato'] if privacy else f"{d['nome']} {d['cognome']}"
                cd = d['credito_debito'] or 0

                if cd >= 0:
                    cd_f = credit_alt_fmt if is_alt else credit_fmt
                else:
                    cd_f = debit_alt_fmt if is_alt else debit_fmt

                ws_mese.write(r, 0, d['commessa'], cf)
                ws_mese.write(r, 1, d['scuola'], cf)
                ws_mese.write(r, 2, utente, cf)
                ws_mese.write(r, 3, d['monte_ore_settimanale'], nf)
                ws_mese.write(r, 4, round(d['media_mensile_60'] or 0, 2), nf)
                ws_mese.write(r, 5, round(d['media_con_assenza_60'] or 0, 2), nf)
                ws_mese.write(r, 6, decimal_to_sessagesimal(d['ore_lavorate_60'] or 0), cf)
                ws_mese.write(r, 7, d['ore_lavorate_100'] or 0, nf)
                ws_mese.write(r, 8, d['imponibile_100'] or 0, mf)
                ws_mese.write(r, 9, d['iva_100'] or 0, mf)
                ws_mese.write(r, 10, d['totale_100'] or 0, mf)
                ws_mese.write(r, 11, cd, cd_f)
                ws_mese.write(r, 12, d['pasti'] or 0, nf)

            # Riga totali
            if dati:
                total_row = 3 + len(dati)
                ore_tot_60 = sum(d['ore_lavorate_60'] or 0 for d in dati)
                ore_tot_100 = sum(d['ore_lavorate_100'] or 0 for d in dati)
                imp_tot = round(ore_tot_100 * TARIFFA, 2)
                iva_tot = round(imp_tot * IVA_PERC, 2)
                tot_tot = round(imp_tot + iva_tot, 2)
                cd_tot = sum(d['credito_debito'] or 0 for d in dati)
                pasti_tot = sum(d['pasti'] or 0 for d in dati)

                ws_mese.write(total_row, 0, 'TOTALE', total_text_fmt)
                for c in range(1, 6):
                    ws_mese.write(total_row, c, '', total_fmt)
                ws_mese.write(total_row, 6, decimal_to_sessagesimal(ore_tot_60), total_fmt)
                ws_mese.write(total_row, 7, ore_tot_100, total_fmt)
                ws_mese.write(total_row, 8, imp_tot, total_money_fmt)
                ws_mese.write(total_row, 9, iva_tot, total_money_fmt)
                ws_mese.write(total_row, 10, tot_tot, total_money_fmt)
                ws_mese.write(total_row, 11, round(cd_tot, 2), total_fmt)
                ws_mese.write(total_row, 12, pasti_tot, total_fmt)

            # Larghezza colonne
            ws_mese.set_column('A:A', 12)
            ws_mese.set_column('B:B', 45)
            ws_mese.set_column('C:C', 22)
            ws_mese.set_column('D:M', 12)

            # Freeze header
            ws_mese.freeze_panes(3, 0)

    output.seek(0)

    filename = f"rendicontazione_annuale_{anno_scolastico}"
    if commessa:
        filename += f"_{commessa.replace(' ', '_')}"
    filename += ".xlsx"

    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename
    )


# ==================== REPORT MUNICIPALE E DIPARTIMENTALE ====================

def classifica_livello_scolastico(scuola_nome):
    """
    Classifica una scuola per livello scolastico in base al nome.
    Ritorna una tupla (livello, ordine) per ordinamento.

    Esempi:
    - /IC ANAGNI/Infanzia/ANAGNI - Via Anagni, 48 → Infanzia Statale (dentro IC)
    - /INFANZIA COMUNALE/TOTI - Via del Pigneto, 104 → Infanzia Capitolina
    - /IC ANAGNI/Primaria/FIUGGI - Via Fiuggi, 18 → Primaria e Secondaria I° Statale
    - /IC ANAGNI/Secondaria/TONIOLO - Via Anagni, 46 → Primaria e Secondaria I° Statale
    """
    nome_upper = scuola_nome.upper()

    # 1. Infanzia Capitolina = contiene "INFANZIA COMUNALE"
    if 'INFANZIA COMUNALE' in nome_upper:
        return ('Infanzia Capitolina', 1)

    # 2. Scuole Paritarie = contiene "PARITARI"
    if 'PARITARI' in nome_upper:
        if 'INFANZIA' in nome_upper:
            return ('Infanzia Paritaria', 3)
        else:
            return ('Primaria e Secondaria I° Paritaria', 5)

    # 3. Scuole dentro IC = contiene "/IC " all'inizio del path
    # Es: /IC ANAGNI/Infanzia/... o /IC ANAGNI/Primaria/...
    if '/IC ' in nome_upper:
        # Cerca il tipo di scuola nel nome
        if 'INFANZIA' in nome_upper:
            return ('Infanzia Statale', 2)
        elif 'PRIMARIA' in nome_upper or 'SECONDARIA' in nome_upper:
            return ('Primaria e Secondaria I° Statale', 4)
        else:
            # IC senza tipo specifico → default primaria/secondaria
            return ('Primaria e Secondaria I° Statale', 4)

    # 4. Altre scuole con "INFANZIA" ma senza "COMUNALE" e senza "IC"
    # → probabilmente capitoline non etichettate correttamente
    if 'INFANZIA' in nome_upper:
        return ('Infanzia Capitolina', 1)

    # 5. Default finale
    return ('Primaria e Secondaria I° Statale', 4)


@export_bp.route('/api/export/municipale/<int:anno>/<int:mese>')
def api_export_municipale(anno, mese):
    """Esporta Riepilogo Municipale - Report per il Municipio"""
    commessa = request.args.get('commessa')

    dati = db.get_rendicontazione_completa(anno, mese, commessa)
    totali_scuola = db.get_totali_per_scuola(anno, mese, commessa)

    # Costanti
    TARIFFA = config.TARIFFA_ORARIA
    IVA_PERC = config.IVA_PERCENTUALE

    # Determina anno scolastico
    if mese >= 9:
        anno_scolastico = f"{anno}/{anno+1}"
    else:
        anno_scolastico = f"{anno-1}/{anno}"

    # Calcola totali generali
    totale_generale = {
        'num_utenti': len(dati),
        'utenti_lista_attesa': sum(1 for d in dati if d.get('lista_attesa')),
        'ore_previste': sum(d['media_con_assenza_60'] or 0 for d in dati),
        'ore_erogate_60': sum(d['ore_lavorate_60'] or 0 for d in dati),
        'ore_erogate_100': sum(d['ore_lavorate_100'] or 0 for d in dati),
    }

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        workbook = writer.book
        s = get_excel_brand_styles(workbook)

        # ========== FOGLIO 1: RIEPILOGO MUNICIPALE ==========
        ws = workbook.add_worksheet('Riepilogo Municipale')
        ws.hide_gridlines(2)

        N_COLS = 7  # 7 colonne tabella scuole

        # Titolo
        ws.set_row(0, 36)
        ws.merge_range(0, 0, 0, N_COLS - 1,
            "RIEPILOGO MUNICIPALE", s['title'])

        # Sottotitolo (mese / anno scolastico / commessa)
        sottotitolo = f"{MESI_NOME[mese].upper()} {anno}   ·   A.S. {anno_scolastico}"
        if commessa:
            sottotitolo += f"   ·   Commessa: {commessa}"
        ws.set_row(1, 22)
        ws.merge_range(1, 0, 1, N_COLS - 1, sottotitolo, s['subtitle'])

        # Riga vuota spaziatrice
        ws.set_row(2, 8)

        # Sezione "DETTAGLIO PER SCUOLA"
        ws.set_row(3, 22)
        ws.merge_range(3, 0, 3, N_COLS - 1, '  DETTAGLIO PER ISTITUTO/SCUOLA', s['section'])

        # Headers tabella
        row = 4
        headers = [
            'Istituto Comprensivo / Scuola',
            "Ore erogate (60')",
            'Tariffa',
            "Totale (60')",
            "Ore (100')",
            'Tariffa',
            'Totale €'
        ]
        ws.set_row(row, 32)
        for col, h in enumerate(headers):
            ws.write(row, col, h, s['header'])

        # Freeze panes sotto l'header
        ws.freeze_panes(row + 1, 0)

        # Dati per scuola con righe alternate
        row = 5
        tot_ore_60 = 0
        tot_ore_100 = 0
        tot_importo_60 = 0
        tot_importo_100 = 0

        for idx, t in enumerate(totali_scuola):
            alt = (idx % 2 == 1)
            ore_60 = t['ore_lavorate_60'] or 0
            ore_100 = t['ore_lavorate_100'] or 0
            importo_60 = ore_60 * TARIFFA
            importo_100 = ore_100 * TARIFFA

            cell_f = s['cell_alt'] if alt else s['cell']
            cell_c_f = s['cell_center_alt'] if alt else s['cell_center']
            num_f = s['number_alt'] if alt else s['number']
            money_f = s['money_alt'] if alt else s['money']

            ws.write(row, 0, t['scuola'], cell_f)
            ws.write(row, 1, decimal_to_sessagesimal(ore_60), cell_c_f)
            ws.write(row, 2, TARIFFA, money_f)
            ws.write(row, 3, importo_60, money_f)
            ws.write(row, 4, ore_100, num_f)
            ws.write(row, 5, TARIFFA, money_f)
            ws.write(row, 6, importo_100, money_f)

            tot_ore_60 += ore_60
            tot_ore_100 += ore_100
            tot_importo_60 += importo_60
            tot_importo_100 += importo_100
            row += 1

        # Riga totale
        ws.set_row(row, 24)
        ws.write(row, 0, 'TOTALE', s['total_label'])
        ws.write(row, 1, decimal_to_sessagesimal(tot_ore_60), s['total_cell'])
        ws.write(row, 2, TARIFFA, s['total_money'])
        ws.write(row, 3, tot_importo_60, s['total_money'])
        ws.write(row, 4, tot_ore_100, s['total_number'])
        ws.write(row, 5, TARIFFA, s['total_money'])
        ws.write(row, 6, tot_importo_100, s['total_money'])

        # Calcoli finali — Box di sintesi fatturazione
        imponibile = tot_importo_100
        iva = imponibile * IVA_PERC
        totale_fatturare = imponibile + iva

        row += 2
        ws.set_row(row, 22)
        ws.merge_range(row, 0, row, N_COLS - 1, '  RIEPILOGO FATTURAZIONE', s['section'])
        row += 1
        ws.write(row, 5, 'Imponibile', s['summary_label'])
        ws.write(row, 6, imponibile, s['summary_money'])
        row += 1
        ws.write(row, 5, f"IVA {int(IVA_PERC * 100)}%", s['summary_label'])
        ws.write(row, 6, iva, s['summary_money'])
        row += 1
        ws.set_row(row, 24)
        ws.write(row, 5, 'IMPORTO DA FATTURARE', s['grand_label'])
        ws.write(row, 6, totale_fatturare, s['grand_money'])

        # ========== Sezione Riepilogativo per Lista di Attesa ==========
        row += 3

        # Liste di attesa distinte ordinate cronologicamente
        liste_attesa = get_liste_attesa_ordinate(dati, anno, mese)

        # Suddividi utenti
        utenti_non_lista = [d for d in dati if not d.get('lista_attesa')]
        utenti_in_lista_totali = [d for d in dati if d.get('lista_attesa')]
        utenti_per_lista = {l['valore']: [d for d in dati if (d.get('lista_attesa') or '').strip() == l['valore']]
                            for l in liste_attesa}

        def _conta_con_ore(lst):
            return sum(1 for d in lst if (d['ore_lavorate_60'] or 0) > 0)

        def _somma_ore_100(lst):
            return sum(d['ore_lavorate_100'] or 0 for d in lst)

        n_col_riepilogo = 4 + len(liste_attesa)
        section_end_col = max(n_col_riepilogo - 1, N_COLS - 1)

        ws.set_row(row, 22)
        ws.merge_range(row, 0, row, section_end_col, '  RIEPILOGATIVO PER LISTA DI ATTESA', s['section'])
        row += 1

        # Header
        riepilogo_headers = ['Indicatore', 'Utenti serviti totali', 'Non in lista attesa', 'Di cui in lista di attesa']
        for l in liste_attesa:
            riepilogo_headers.append(l['label'])
        ws.set_row(row, 32)
        for col, h in enumerate(riepilogo_headers):
            ws.write(row, col, h, s['header'])

        # Riga 1: Alunni/ore/importo (numero utenti)
        row += 1
        ws.write(row, 0, 'Alunni assistiti (totale)', s['cell'])
        ws.write(row, 1, totale_generale['num_utenti'], s['integer'])
        ws.write(row, 2, len(utenti_non_lista), s['integer'])
        ws.write(row, 3, len(utenti_in_lista_totali), s['integer'])
        for i, l in enumerate(liste_attesa):
            ws.write(row, 4 + i, len(utenti_per_lista[l['valore']]), s['integer'])

        # Riga 2: Alunni effettivamente assistiti
        row += 1
        ws.write(row, 0, 'Alunni effettivamente assistiti nel mese', s['cell_alt'])
        ws.write(row, 1, _conta_con_ore(dati), s['integer_alt'])
        ws.write(row, 2, _conta_con_ore(utenti_non_lista), s['integer_alt'])
        ws.write(row, 3, _conta_con_ore(utenti_in_lista_totali), s['integer_alt'])
        for i, l in enumerate(liste_attesa):
            ws.write(row, 4 + i, _conta_con_ore(utenti_per_lista[l['valore']]), s['integer_alt'])

        # Riga 3: Ore erogate (100')
        row += 1
        ore_100_non_lista = _somma_ore_100(utenti_non_lista)
        ore_100_in_lista = _somma_ore_100(utenti_in_lista_totali)
        ws.write(row, 0, "Ore effettivamente erogate (al netto dell'11%)", s['cell'])
        ws.write(row, 1, tot_ore_100, s['number'])
        ws.write(row, 2, ore_100_non_lista, s['number'])
        ws.write(row, 3, ore_100_in_lista, s['number'])
        for i, l in enumerate(liste_attesa):
            ws.write(row, 4 + i, _somma_ore_100(utenti_per_lista[l['valore']]), s['number'])

        # Riga 4: Importo (imponibile + IVA)
        row += 1
        importo_non_lista = ore_100_non_lista * TARIFFA
        importo_in_lista = ore_100_in_lista * TARIFFA
        totale_non_lista = importo_non_lista * (1 + IVA_PERC)
        totale_in_lista = importo_in_lista * (1 + IVA_PERC)
        ws.write(row, 0, 'Importo effettivamente erogato (IVA inclusa)', s['cell_alt'])
        ws.write(row, 1, totale_fatturare, s['money_alt'])
        ws.write(row, 2, totale_non_lista, s['money_alt'])
        ws.write(row, 3, totale_in_lista, s['money_alt'])
        for i, l in enumerate(liste_attesa):
            ore_lista = _somma_ore_100(utenti_per_lista[l['valore']])
            tot_lista = ore_lista * TARIFFA * (1 + IVA_PERC)
            ws.write(row, 4 + i, tot_lista, s['money_alt'])

        # Footer informativo
        row += 2
        ws.write(row, 0, f"Documento generato il {datetime.now().strftime('%d/%m/%Y %H:%M')}", s['info'])

        # Larghezze colonne
        ws.set_column('A:A', 52)
        ws.set_column(1, max(N_COLS - 1, n_col_riepilogo - 1), 18)

        # Margini per stampa
        ws.set_margins(left=0.5, right=0.5, top=0.5, bottom=0.5)
        ws.set_landscape()
        ws.fit_to_pages(1, 0)
        ws.repeat_rows(0, 1)

        # ========== FOGLIO 2: DETTAGLIO UTENTI ==========
        ws_utenti = workbook.add_worksheet('Dettaglio Utenti')
        ws_utenti.hide_gridlines(2)

        N_COLS_U = 7

        # Titolo
        ws_utenti.set_row(0, 36)
        ws_utenti.merge_range(0, 0, 0, N_COLS_U - 1, 'DETTAGLIO UTENTI', s['title'])

        # Sottotitolo
        ws_utenti.set_row(1, 22)
        sub_u = f"{MESI_NOME[mese].upper()} {anno}   ·   A.S. {anno_scolastico}"
        if commessa:
            sub_u += f"   ·   Commessa: {commessa}"
        ws_utenti.merge_range(1, 0, 1, N_COLS_U - 1, sub_u, s['subtitle'])

        ws_utenti.set_row(2, 8)

        # Headers tabella utenti
        utenti_headers = ['Scuola', 'Utente', 'Monte Ore', "Ore Erogate (60')", "Ore (100')", 'Totale €', 'Lista Attesa']
        row_u = 3
        ws_utenti.set_row(row_u, 32)
        for col, h in enumerate(utenti_headers):
            ws_utenti.write(row_u, col, h, s['header'])

        ws_utenti.freeze_panes(row_u + 1, 0)

        # Dati utenti con righe alternate
        row_u = 4
        for idx, d in enumerate(dati):
            alt = (idx % 2 == 1)
            cell_f = s['cell_alt'] if alt else s['cell']
            cell_c_f = s['cell_center_alt'] if alt else s['cell_center']
            num_f = s['number_alt'] if alt else s['number']
            money_f = s['money_alt'] if alt else s['money']

            ws_utenti.write(row_u, 0, d['scuola'], cell_f)
            ws_utenti.write(row_u, 1, f"{d['nome']} {d['cognome']}", cell_f)
            ws_utenti.write(row_u, 2, d['monte_ore_settimanale'], num_f)
            ws_utenti.write(row_u, 3, decimal_to_sessagesimal(d['ore_lavorate_60'] or 0), cell_c_f)
            ws_utenti.write(row_u, 4, d['ore_lavorate_100'] or 0, num_f)
            ws_utenti.write(row_u, 5, d['totale_100'] or 0, money_f)
            ws_utenti.write(row_u, 6, d.get('lista_attesa') or '—', cell_c_f)
            row_u += 1

        # Footer informativo
        row_u += 1
        ws_utenti.write(row_u, 0, f"Documento generato il {datetime.now().strftime('%d/%m/%Y %H:%M')}", s['info'])

        # Larghezza colonne foglio utenti
        ws_utenti.set_column('A:A', 50)
        ws_utenti.set_column('B:B', 26)
        ws_utenti.set_column('C:C', 12)
        ws_utenti.set_column('D:F', 16)
        ws_utenti.set_column('G:G', 14)

        # Margini per stampa
        ws_utenti.set_margins(left=0.5, right=0.5, top=0.5, bottom=0.5)
        ws_utenti.set_landscape()
        ws_utenti.fit_to_pages(1, 0)
        ws_utenti.repeat_rows(0, 3)

    output.seek(0)

    filename = f"Riepilogo_Municipale_{MESI_NOME[mese]}_{anno}"
    if commessa:
        filename += f"_{commessa.replace(' ', '_')}"
    filename += ".xlsx"

    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename
    )


@export_bp.route('/api/export/dipartimentale/<int:anno>/<int:mese>')
def api_export_dipartimentale(anno, mese):
    """Esporta Monitoraggio Dipartimentale - Report per livello scolastico"""
    commessa = request.args.get('commessa')

    dati = db.get_rendicontazione_completa(anno, mese, commessa)

    # Costanti (stesse del riepilogo municipale per coerenza)
    TARIFFA = config.TARIFFA_ORARIA
    IVA_PERC = config.IVA_PERCENTUALE

    # Determina anno scolastico
    if mese >= 9:
        anno_scolastico = f"{anno}/{anno+1}"
    else:
        anno_scolastico = f"{anno-1}/{anno}"

    # Raggruppa per livello scolastico
    livelli = {}
    for d in dati:
        livello, ordine = classifica_livello_scolastico(d['scuola'])

        if livello not in livelli:
            livelli[livello] = {
                'ordine': ordine,
                'n_utenti': 0,
                'ore_richieste': 0,  # media -11%
                'ore_erogate': 0,    # ore lavorate 100'
                'importo_impegnato': 0,  # media -11% × tariffa con IVA
                'importo_liquidato': 0   # ore erogate × tariffa con IVA
            }

        ore_richieste = d['media_con_assenza_60'] or 0
        ore_erogate = d['ore_lavorate_100'] or 0

        livelli[livello]['n_utenti'] += 1
        livelli[livello]['ore_richieste'] += ore_richieste
        livelli[livello]['ore_erogate'] += ore_erogate
        livelli[livello]['importo_impegnato'] += ore_richieste * TARIFFA * (1 + IVA_PERC)
        livelli[livello]['importo_liquidato'] += ore_erogate * TARIFFA * (1 + IVA_PERC)

    # Ordina per ordine predefinito
    livelli_ordinati = sorted(livelli.items(), key=lambda x: x[1]['ordine'])

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        workbook = writer.book
        s = get_excel_brand_styles(workbook)

        # Foglio
        ws = workbook.add_worksheet('Monitoraggio Dipartimentale')
        ws.hide_gridlines(2)

        N_COLS = 6

        # Titolo
        ws.set_row(0, 36)
        ws.merge_range(0, 0, 0, N_COLS - 1, 'MONITORAGGIO DIPARTIMENTALE', s['title'])

        # Sottotitolo
        sub = f"{MESI_NOME[mese].upper()} {anno}   ·   A.S. {anno_scolastico}"
        if commessa:
            sub += f"   ·   Commessa: {commessa}"
        ws.set_row(1, 22)
        ws.merge_range(1, 0, 1, N_COLS - 1, sub, s['subtitle'])

        # Spazio
        ws.set_row(2, 8)

        # Sezione
        ws.set_row(3, 22)
        ws.merge_range(3, 0, 3, N_COLS - 1, '  RIEPILOGO PER LIVELLO SCOLASTICO', s['section'])

        # ========== TABELLA DATI ==========
        # Headers
        row = 4
        headers = [
            'Livello Scolastico',
            'N. Utenti',
            'Ore richieste',
            'Ore Erogate',
            'Importo Impegnato',
            'Importo Liquidato'
        ]

        ws.set_row(row, 36)
        for col, h in enumerate(headers):
            ws.write(row, col, h, s['header'])

        ws.freeze_panes(row + 1, 0)

        # Dati per livello (righe alternate)
        row = 5
        tot_utenti = 0
        tot_ore_richieste = 0
        tot_ore_erogate = 0
        tot_impegnato = 0
        tot_liquidato = 0

        for idx, (livello, stats) in enumerate(livelli_ordinati):
            alt = (idx % 2 == 1)
            cell_f = s['cell_alt'] if alt else s['cell']
            num_f = s['number_alt'] if alt else s['number']
            int_f = s['integer_alt'] if alt else s['integer']
            money_f = s['money_alt'] if alt else s['money']

            ws.write(row, 0, livello, cell_f)
            ws.write(row, 1, stats['n_utenti'], int_f)
            ws.write(row, 2, stats['ore_richieste'], num_f)
            ws.write(row, 3, stats['ore_erogate'], num_f)
            ws.write(row, 4, stats['importo_impegnato'], money_f)
            ws.write(row, 5, stats['importo_liquidato'], money_f)

            tot_utenti += stats['n_utenti']
            tot_ore_richieste += stats['ore_richieste']
            tot_ore_erogate += stats['ore_erogate']
            tot_impegnato += stats['importo_impegnato']
            tot_liquidato += stats['importo_liquidato']

            row += 1

        # Riga TOTALE
        ws.set_row(row, 26)
        ws.write(row, 0, 'TOTALE', s['total_label'])
        ws.write(row, 1, tot_utenti, s['total_integer'])
        ws.write(row, 2, tot_ore_richieste, s['total_number'])
        ws.write(row, 3, tot_ore_erogate, s['total_number'])
        ws.write(row, 4, tot_impegnato, s['total_money'])
        ws.write(row, 5, tot_liquidato, s['total_money'])

        # Footer informativo
        row += 2
        ws.write(row, 0, f"Documento generato il {datetime.now().strftime('%d/%m/%Y %H:%M')}", s['info'])

        # Larghezza colonne
        ws.set_column('A:A', 30)
        ws.set_column('B:B', 12)
        ws.set_column('C:D', 16)
        ws.set_column('E:F', 20)

        # Stampa
        ws.set_margins(left=0.5, right=0.5, top=0.5, bottom=0.5)
        ws.set_landscape()
        ws.fit_to_pages(1, 0)
        ws.repeat_rows(0, 1)

    output.seek(0)

    filename = f"Monitoraggio_Dipartimentale_{MESI_NOME[mese]}_{anno}"
    if commessa:
        filename += f"_{commessa.replace(' ', '_')}"
    filename += ".xlsx"

    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
        as_attachment=True,
        download_name=filename
    )


# ==================== STATISTICHE ====================

@export_bp.route('/api/export/word/<int:anno>/<int:mese>')
def api_export_word(anno, mese):
    """Genera un documento Word con relazione sull'andamento del servizio mensile"""
    commessa = request.args.get('commessa')

    dati = db.get_rendicontazione_completa(anno, mese, commessa)
    totali_scuola = db.get_totali_per_scuola(anno, mese, commessa)

    # Costanti
    TARIFFA = config.TARIFFA_ORARIA
    IVA_PERC = config.IVA_PERCENTUALE

    # Calcola totali
    ore_totali_60 = sum(d['ore_lavorate_60'] or 0 for d in dati)
    ore_totali_100 = sum(d['ore_lavorate_100'] or 0 for d in dati)
    ore_previste = sum(d['media_con_assenza_60'] or 0 for d in dati)
    pasti_totali = sum(d['pasti'] or 0 for d in dati)
    credito_debito = sum(d['credito_debito'] or 0 for d in dati)

    # Calcolo fatturazione
    imponibile_totale, iva_totale, totale_lordo = config.calcola_fatturazione(ore_totali_100)

    # Percentuale completamento
    perc_completamento = (ore_totali_60 / ore_previste * 100) if ore_previste > 0 else 0

    # Utenti con ore e in lista attesa
    utenti_con_ore = sum(1 for d in dati if (d['ore_lavorate_60'] or 0) > 0)
    utenti_lista_attesa = sum(1 for d in dati if d.get('lista_attesa'))

    # Determina anno scolastico
    if mese >= 9:
        anno_scolastico = f"{anno}/{anno+1}"
    else:
        anno_scolastico = f"{anno-1}/{anno}"

    # Crea documento Word
    doc = Document()

    # Stile di default del documento
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)

    # Imposta margini
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Footer con numero pagina
    add_word_page_number_footer(doc)

    # ========== INTESTAZIONE BRANDED ==========
    # Barra colorata superiore (titolo con sfondo indigo)
    title_table = doc.add_table(rows=1, cols=1)
    title_cell = title_table.rows[0].cells[0]
    _style_word_cell_bg(title_cell, '4F46E5')
    title_p = title_cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run('RELAZIONE SUL SERVIZIO OEPAC')
    title_run.bold = True
    title_run.font.size = Pt(22)
    title_run.font.color.rgb = RGBColor.from_string('FFFFFF')
    title_run.font.name = 'Calibri'

    # Sottotitolo con sfondo dark
    sub_table = doc.add_table(rows=1, cols=1)
    sub_cell = sub_table.rows[0].cells[0]
    _style_word_cell_bg(sub_cell, '1E293B')
    sub_p = sub_cell.paragraphs[0]
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(f'{MESI_NOME[mese].upper()} {anno}   ·   A.S. {anno_scolastico}')
    sub_run.bold = True
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor.from_string('FFFFFF')
    sub_run.font.name = 'Calibri'

    if commessa:
        comm_para = doc.add_paragraph()
        comm_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = comm_para.add_run(f'Commessa: {commessa}')
        run.font.size = Pt(11)
        run.italic = True
        run.font.color.rgb = RGBColor.from_string('64748B')

    # Data di emissione
    data_em = doc.add_paragraph()
    data_em.alignment = WD_ALIGN_PARAGRAPH.CENTER
    de_run = data_em.add_run(f"Emesso il {datetime.now().strftime('%d/%m/%Y')}")
    de_run.font.size = Pt(10)
    de_run.italic = True
    de_run.font.color.rgb = RGBColor.from_string('64748B')

    doc.add_paragraph()

    # Sezione 1: Panoramica Generale
    h1 = doc.add_heading('1. PANORAMICA GENERALE', level=1)
    for run in h1.runs:
        run.font.color.rgb = RGBColor.from_string('4F46E5')

    intro = doc.add_paragraph()
    intro.add_run(f"Nel mese di {MESI_NOME[mese]} {anno}, relativo all'anno scolastico {anno_scolastico}, "
                  f"il servizio OEPAC ha assistito un totale di ").bold = False
    intro.add_run(f"{len(dati)} utenti").bold = True
    intro.add_run(" distribuiti su ")
    intro.add_run(f"{len(totali_scuola)} scuole").bold = True
    intro.add_run(".")

    if utenti_lista_attesa > 0:
        lista_para = doc.add_paragraph()
        lista_para.add_run(f"Di questi, {utenti_lista_attesa} utenti risultano in lista di attesa.").italic = True

    # Sezione 2: Dati Quantitativi
    h2 = doc.add_heading('2. DATI QUANTITATIVI', level=1)
    for run in h2.runs:
        run.font.color.rgb = RGBColor.from_string('4F46E5')

    # Tabella KPI
    table_kpi = doc.add_table(rows=1, cols=2)
    table_kpi.style = 'Table Grid'
    table_kpi.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_kpi.rows[0].cells
    hdr_cells[0].text = 'Indicatore'
    hdr_cells[1].text = 'Valore'

    kpi_data = [
        ('Utenti totali assistiti', str(len(dati))),
        ('Utenti effettivamente serviti nel mese', str(utenti_con_ore)),
        ('Utenti in lista di attesa', str(utenti_lista_attesa)),
        ('Ore previste (Media -11%)', f'{ore_previste:.2f}'),
        ('Ore effettivamente erogate', f'{ore_totali_60:.2f}'),
        ('Percentuale completamento', f'{perc_completamento:.1f}%'),
        ('Pasti erogati', str(pasti_totali)),
        ('Credito/Debito ore', f'{credito_debito:+.2f}'),
    ]

    for label, value in kpi_data:
        row_cells = table_kpi.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    style_word_table_header(table_kpi)
    style_word_table_alternating_rows(table_kpi)

    doc.add_paragraph()

    # Sezione 3: Dati Economici
    h3 = doc.add_heading('3. DATI ECONOMICI', level=1)
    for run in h3.runs:
        run.font.color.rgb = RGBColor.from_string('4F46E5')

    econ_para = doc.add_paragraph()
    econ_para.add_run(f"Sulla base delle ore erogate nel mese di {MESI_NOME[mese]}, "
                      f"applicando la tariffa oraria di € {TARIFFA:.2f} (esclusa IVA), "
                      f"si riportano i seguenti dati economici:")

    table_econ = doc.add_table(rows=1, cols=2)
    table_econ.style = 'Table Grid'
    table_econ.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table_econ.rows[0].cells
    hdr_cells[0].text = 'Voce'
    hdr_cells[1].text = 'Importo'

    econ_data = [
        ('Ore erogate (centesimali)', f'{ore_totali_100:.2f}'),
        ('Imponibile', f'€ {imponibile_totale:,.2f}'),
        (f'IVA {int(IVA_PERC * 100)}%', f'€ {iva_totale:,.2f}'),
        ('TOTALE DA FATTURARE', f'€ {totale_lordo:,.2f}'),
    ]

    for label, value in econ_data:
        row_cells = table_econ.add_row().cells
        row_cells[0].text = label
        row_cells[1].text = value

    style_word_table_header(table_econ)
    style_word_table_alternating_rows(table_econ)
    style_word_table_total_row(table_econ, bg_color='4F46E5')

    doc.add_paragraph()

    # Sezione 4: Riepilogativo per Lista di Attesa
    h4 = doc.add_heading('4. RIEPILOGATIVO PER LISTA DI ATTESA', level=1)
    for run in h4.runs:
        run.font.color.rgb = RGBColor.from_string('4F46E5')

    liste_attesa = get_liste_attesa_ordinate(dati, anno, mese)
    utenti_non_lista_rel = [d for d in dati if not d.get('lista_attesa')]
    utenti_in_lista_rel = [d for d in dati if d.get('lista_attesa')]
    utenti_per_lista_rel = {l['valore']: [d for d in dati if (d.get('lista_attesa') or '').strip() == l['valore']]
                            for l in liste_attesa}

    def _conta_con_ore_rel(lst):
        return sum(1 for d in lst if (d['ore_lavorate_60'] or 0) > 0)

    def _somma_ore_100_rel(lst):
        return sum(d['ore_lavorate_100'] or 0 for d in lst)

    # Tabella riepilogativo: 4 righe x (4 + N liste) colonne
    n_cols_riep = 4 + len(liste_attesa)
    table_riep = doc.add_table(rows=1, cols=n_cols_riep)
    table_riep.style = 'Table Grid'
    table_riep.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Header
    hdr = table_riep.rows[0].cells
    hdr[0].text = 'Indicatore'
    hdr[1].text = 'Utenti serviti totali'
    hdr[2].text = 'Non in lista attesa'
    hdr[3].text = 'Di cui in lista di attesa'
    for i, l in enumerate(liste_attesa):
        hdr[4 + i].text = l['label']

    # Riga 1: Alunni
    r1 = table_riep.add_row().cells
    r1[0].text = 'Alunni assistiti (totale)'
    r1[1].text = str(len(dati))
    r1[2].text = str(len(utenti_non_lista_rel))
    r1[3].text = str(len(utenti_in_lista_rel))
    for i, l in enumerate(liste_attesa):
        r1[4 + i].text = str(len(utenti_per_lista_rel[l['valore']]))

    # Riga 2: Alunni assistiti
    r2 = table_riep.add_row().cells
    r2[0].text = 'Alunni effettivamente assistiti nel mese'
    r2[1].text = str(_conta_con_ore_rel(dati))
    r2[2].text = str(_conta_con_ore_rel(utenti_non_lista_rel))
    r2[3].text = str(_conta_con_ore_rel(utenti_in_lista_rel))
    for i, l in enumerate(liste_attesa):
        r2[4 + i].text = str(_conta_con_ore_rel(utenti_per_lista_rel[l['valore']]))

    # Riga 3: Ore erogate (100')
    ore_100_non_lista_rel = _somma_ore_100_rel(utenti_non_lista_rel)
    ore_100_in_lista_rel = _somma_ore_100_rel(utenti_in_lista_rel)
    r3 = table_riep.add_row().cells
    r3[0].text = "Ore effettivamente erogate (al netto dell'11%)"
    r3[1].text = f'{ore_totali_100:.2f}'
    r3[2].text = f'{ore_100_non_lista_rel:.2f}'
    r3[3].text = f'{ore_100_in_lista_rel:.2f}'
    for i, l in enumerate(liste_attesa):
        r3[4 + i].text = f'{_somma_ore_100_rel(utenti_per_lista_rel[l["valore"]]):.2f}'

    # Riga 4: Importo (imponibile + IVA)
    importo_non_lista_rel = ore_100_non_lista_rel * TARIFFA * (1 + IVA_PERC)
    importo_in_lista_rel = ore_100_in_lista_rel * TARIFFA * (1 + IVA_PERC)
    r4 = table_riep.add_row().cells
    r4[0].text = 'Importo erogato (IVA inclusa)'
    r4[1].text = f'€ {totale_lordo:,.2f}'
    r4[2].text = f'€ {importo_non_lista_rel:,.2f}'
    r4[3].text = f'€ {importo_in_lista_rel:,.2f}'
    for i, l in enumerate(liste_attesa):
        ore_l = _somma_ore_100_rel(utenti_per_lista_rel[l['valore']])
        imp_l = ore_l * TARIFFA * (1 + IVA_PERC)
        r4[4 + i].text = f'€ {imp_l:,.2f}'

    style_word_table_header(table_riep)
    style_word_table_alternating_rows(table_riep)

    doc.add_paragraph()

    # Sezione 5: Distribuzione per Scuola
    h5 = doc.add_heading('5. DISTRIBUZIONE PER SCUOLA', level=1)
    for run in h5.runs:
        run.font.color.rgb = RGBColor.from_string('4F46E5')

    if totali_scuola:
        distr_para = doc.add_paragraph()
        distr_para.add_run("Di seguito il riepilogo delle ore erogate suddivise per scuola:")

        table_scuole = doc.add_table(rows=1, cols=5)
        table_scuole.style = 'Table Grid'
        table_scuole.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table_scuole.rows[0].cells
        headers = ['Scuola', 'Utenti', 'Ore Erogate', 'Imponibile', 'Totale']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h

        # Totalizzatori per la riga TOTALE
        tot_utenti_s = 0
        tot_ore_s = 0.0
        tot_imp_s = 0.0
        tot_tot_s = 0.0

        for t in totali_scuola:
            row_cells = table_scuole.add_row().cells
            nome_scuola = t['scuola'][:50] + '...' if len(t['scuola']) > 50 else t['scuola']
            row_cells[0].text = nome_scuola
            row_cells[1].text = str(t['num_utenti'])
            row_cells[2].text = f"{t['ore_lavorate_60']:.2f}"
            row_cells[3].text = f"€ {t['imponibile_100']:,.2f}"
            row_cells[4].text = f"€ {t['totale_100']:,.2f}"

            tot_utenti_s += t['num_utenti'] or 0
            tot_ore_s += t['ore_lavorate_60'] or 0
            tot_imp_s += t['imponibile_100'] or 0
            tot_tot_s += t['totale_100'] or 0

        # Riga TOTALE
        tot_cells = table_scuole.add_row().cells
        tot_cells[0].text = 'TOTALE'
        tot_cells[1].text = str(tot_utenti_s)
        tot_cells[2].text = f"{tot_ore_s:.2f}"
        tot_cells[3].text = f"€ {tot_imp_s:,.2f}"
        tot_cells[4].text = f"€ {tot_tot_s:,.2f}"

        style_word_table_header(table_scuole)
        style_word_table_alternating_rows(table_scuole)
        style_word_table_total_row(table_scuole)

    doc.add_paragraph()

    # Sezione 6: Analisi e Osservazioni
    h6 = doc.add_heading('6. ANALISI E OSSERVAZIONI', level=1)
    for run in h6.runs:
        run.font.color.rgb = RGBColor.from_string('4F46E5')

    # Analisi automatica basata sui dati
    if perc_completamento >= 95:
        analisi = f"Il servizio ha raggiunto un ottimo livello di completamento ({perc_completamento:.1f}%), " \
                  f"superando il 95% delle ore previste."
    elif perc_completamento >= 80:
        analisi = f"Il servizio ha raggiunto un buon livello di completamento ({perc_completamento:.1f}%), " \
                  f"erogando oltre l'80% delle ore previste."
    elif perc_completamento >= 60:
        analisi = f"Il servizio ha raggiunto un livello di completamento nella media ({perc_completamento:.1f}%). " \
                  f"Si suggerisce di verificare eventuali criticità."
    else:
        analisi = f"Il livello di completamento ({perc_completamento:.1f}%) risulta inferiore alle aspettative. " \
                  f"Si raccomanda un'analisi approfondita delle cause."

    doc.add_paragraph(analisi)

    if credito_debito > 0:
        cred_para = doc.add_paragraph()
        cred_para.add_run(f"Il saldo credito/debito ore è positivo (+{credito_debito:.2f} ore), "
                         f"indicando ore non ancora erogate rispetto alle previste.")
    elif credito_debito < 0:
        cred_para = doc.add_paragraph()
        cred_para.add_run(f"Il saldo credito/debito ore è negativo ({credito_debito:.2f} ore), "
                         f"indicando ore erogate in eccesso rispetto alle previste.")

    # Sottosezione: Utenti con tasso di erogazione < 50%
    utenti_bassa_erogazione = []
    for d in dati:
        ore_erogate = d['ore_lavorate_60'] or 0
        ore_previste = d['media_con_assenza_60'] or 0
        if ore_previste > 0:
            tasso = (ore_erogate / ore_previste) * 100
            if tasso < 50:
                utenti_bassa_erogazione.append({
                    'nome': f"{d['nome']} {d['cognome']}",
                    'scuola': d['scuola'],
                    'ore_previste': ore_previste,
                    'ore_erogate': ore_erogate,
                    'tasso': tasso
                })

    if utenti_bassa_erogazione:
        doc.add_paragraph()
        h61 = doc.add_heading('6.1 Utenti con tasso di erogazione inferiore al 50%', level=2)
        for run in h61.runs:
            run.font.color.rgb = RGBColor.from_string('EF4444')

        alert_para = doc.add_paragraph()
        alert_para.add_run(f"Si segnalano {len(utenti_bassa_erogazione)} utenti con un tasso di erogazione "
                          f"inferiore al 50% rispetto alle ore previste:")

        table_alert = doc.add_table(rows=1, cols=5)
        table_alert.style = 'Table Grid'
        table_alert.alignment = WD_TABLE_ALIGNMENT.CENTER
        hdr_cells = table_alert.rows[0].cells
        headers_alert = ['Utente', 'Scuola', 'Ore Previste', 'Ore Erogate', 'Tasso']
        for i, h in enumerate(headers_alert):
            hdr_cells[i].text = h

        # Ordina per tasso crescente (i più critici prima)
        utenti_bassa_erogazione.sort(key=lambda x: x['tasso'])

        for u in utenti_bassa_erogazione:
            row_cells = table_alert.add_row().cells
            row_cells[0].text = u['nome']
            nome_scuola = u['scuola'][:40] + '...' if len(u['scuola']) > 40 else u['scuola']
            row_cells[1].text = nome_scuola
            row_cells[2].text = f"{u['ore_previste']:.2f}"
            row_cells[3].text = f"{u['ore_erogate']:.2f}"
            row_cells[4].text = f"{u['tasso']:.1f}%"

        style_word_table_header(table_alert, bg_color='EF4444')
        style_word_table_alternating_rows(table_alert, color_alt='FEE2E2')

    doc.add_paragraph()

    # Sezione 7: Conclusioni
    h7 = doc.add_heading('7. CONCLUSIONI', level=1)
    for run in h7.runs:
        run.font.color.rgb = RGBColor.from_string('4F46E5')

    concl = doc.add_paragraph()
    concl.add_run(f"In sintesi, nel mese di {MESI_NOME[mese]} {anno} il servizio OEPAC ha erogato "
                  f"complessivamente {ore_totali_60:.2f} ore di assistenza a {utenti_con_ore} utenti, "
                  f"per un importo totale da fatturare pari a € {totale_lordo:,.2f}.")

    # Data e firma
    doc.add_paragraph()
    doc.add_paragraph()

    data_para = doc.add_paragraph()
    data_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    dr = data_para.add_run(f"Data: {datetime.now().strftime('%d/%m/%Y')}")
    dr.bold = True

    doc.add_paragraph()

    firma_para = doc.add_paragraph()
    firma_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    fr = firma_para.add_run("Il Responsabile del Servizio")
    fr.bold = True

    firma_line = doc.add_paragraph()
    firma_line.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    firma_line.add_run("_________________________")

    # Salva in buffer
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)

    filename = f"Relazione_OEPAC_{MESI_NOME[mese]}_{anno}"
    if commessa:
        filename += f"_{commessa.replace(' ', '_')}"
    filename += ".docx"

    return send_file(
        output,
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        as_attachment=True,
        download_name=filename
    )
